from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


ROOT = Path("/Users/sasitorn/Project_Stock")
OUT = ROOT / "คู่มือใช้งานเว็บ_Project_Stock_QR.docx"
SHOT = ROOT / "manual_screenshots"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(85, 85, 85)


def add_image(doc, filename, caption):
    path = SHOT / filename
    if path.exists():
        with Image.open(path) as img:
            w, h = img.size
        max_width = 6.25
        max_height = 7.4
        width = max_width
        height = width * h / w
        if height > max_height:
            width = max_height * w / h
        doc.add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_steps(doc, steps):
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "E8EEF5")
        if widths:
            hdr[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    title.font.size = Pt(22)
    title.font.color.rgb = RGBColor(11, 37, 69)

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.188)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25


def add_cover(doc):
    p = doc.add_paragraph()
    p.style = "Title"
    p.add_run("คู่มือใช้งานเว็บ Project Stock QR")
    sub = doc.add_paragraph()
    sub.add_run("แยกขั้นตอนสำหรับ Admin และคนขับ พร้อมตัวอย่างข้อมูลและภาพหน้าจอจากเว็บจริง").italic = True
    meta = [
        ("เว็บไซต์", "https://project-stock-qr.vercel.app/jobs"),
        ("วันที่จัดทำ", "18 พฤษภาคม 2026"),
        ("ระบบที่อ้างอิง", "Store QR Job Transport / DC Bangna"),
        ("Job ตัวอย่างในคู่มือ", "JOB-20260518-152703-000012, ห้อง frt, รถ iu, คนขับ iug"),
    ]
    add_table(doc, ["หัวข้อ", "รายละเอียด"], meta, [Inches(1.65), Inches(4.75)])
    doc.add_paragraph(
        "คู่มือนี้เขียนจากการสำรวจหน้าเว็บจริง โดยใช้ข้อมูลตัวอย่างที่มีอยู่ในระบบเพื่ออธิบาย flow การทำงานตั้งแต่นำเข้า PO จนถึงการสแกนและ monitor งานขนส่ง"
    )


def build():
    doc = Document()
    style_doc(doc)
    add_cover(doc)

    doc.add_heading("ภาพรวมระบบ", level=1)
    doc.add_paragraph(
        "ระบบนี้ใช้จัดการงานขนส่งจาก PO จริง แอดมินจะนำเข้า/เลือก PO แล้วสร้าง Job จากนั้นส่งลิงก์หรือ QR ให้คนขับเปิด Driver Room เพื่อเช็กอิน GPS และสแกนสินค้า ส่วนหน้า Monitor ใช้ดูสถานะโหลดขึ้นรถ ส่งปลายทาง และ alert แบบ realtime"
    )
    add_table(
        doc,
        ["บทบาท", "หน้าที่หลัก", "หน้าที่เกี่ยวข้อง"],
        [
            ("Admin", "นำเข้า PO, เลือกรายการ, สร้าง Job, ดู Monitor, เปิด QR, ตรวจประวัติ", "เมนูฝั่งซ้ายทั้งหมด และปุ่มในตาราง Job"),
            ("คนขับ", "เปิด Driver Room, เช็กอินต้นทาง, สแกนขึ้นรถ, เช็กอินปลายทาง, สแกนส่งของ", "ลิงก์ /driver-room?jobId=... หรือ QR ที่ Admin แสดง"),
        ],
        [Inches(1.2), Inches(3.0), Inches(2.2)],
    )

    doc.add_heading("คู่มือฝั่ง Admin", level=1)

    doc.add_heading("1. เปิดหน้ารายการ Job", level=2)
    add_image(doc, "01-jobs-list.png", "ภาพหน้ารายการ Job ปัจจุบัน มีปุ่ม Monitor, เปิด Driver Room, แสดง QR และลบ Job")
    add_steps(
        doc,
        [
            "เข้าเว็บที่ https://project-stock-qr.vercel.app/jobs",
            "ดูแถบเมนูซ้ายเพื่อเลือกหน้าที่ต้องใช้งาน เช่น นำเข้า PO, PO รอจัดส่ง, สร้าง Job, Monitor Realtime, Driver Room, รายการ Job และประวัติงาน",
            "ในตารางงานทั้งหมด ตรวจคอลัมน์ห้อง Job, Route / PO, Driver, ขึ้นรถ, ส่งแล้ว, Status และ Action",
            "ถ้าต้องติดตามงานให้กด Monitor ถ้าต้องเปิดหน้าคนขับให้กด เปิด Driver Room หรือ แสดง QR สำหรับคนขับ",
        ],
    )

    doc.add_heading("2. นำเข้า PO จากไฟล์", level=2)
    add_image(doc, "02-po-import.png", "ภาพหน้าจอนำเข้า PO จากไฟล์ Excel หรือ CSV")
    add_steps(
        doc,
        [
            "กดเมนู นำเข้า PO",
            "ลากไฟล์มาวางในกรอบอัปโหลด หรือกด เลือกไฟล์ Excel / CSV",
            "ตรวจว่าข้อมูลจากไฟล์ขึ้นในคิว PO รอจัดส่ง",
            "หากข้อมูลไม่ขึ้น ให้ตรวจนามสกุลไฟล์และคอลัมน์หลัก เช่น PO SAP No., Item, Vendor, รหัสวัสดุ, ชื่อวัสดุ และจำนวนสั่งซื้อ",
        ],
    )

    doc.add_heading("3. เลือก PO รอจัดส่ง", level=2)
    add_image(doc, "01-admin-po-queue.png", "ภาพหน้าคิว PO รอจัดส่ง มีช่องค้นหาและ checkbox เลือกรายการ")
    add_steps(
        doc,
        [
            "กดเมนู PO รอจัดส่ง",
            "ค้นหา PO, Vendor, PO Web No. หรือชื่อวัสดุ หากต้องการกรองรายการ",
            "ติ๊ก checkbox หน้ารายการ PO ที่ต้องการสร้าง Job",
            "ตรวจตัวนับ เลือกแล้ว ... รายการ แล้วกด สร้าง Job จากรายการที่เลือก",
        ],
    )
    doc.add_paragraph("ตัวอย่างข้อมูลจากคิว PO:")
    add_table(
        doc,
        ["PO SAP No.", "Item", "Vendor", "รหัสวัสดุ", "สินค้า", "จำนวน"],
        [
            ("8030492670", "1", "บจ.ชวนันท์ คอร์ปอเรชั่น", "80216114", "SMC ชัตเติ้ลวาล์วฟิตติ้งแบบ วันทัช ซีรี่ VR1220F-10", "4"),
            ("8030492670", "2", "บจ.ชวนันท์ คอร์ปอเรชั่น", "80216123", "SMC ตัวควบคุมปริมาณลมพร้อมฟิตติ้งสวมเร็ว AS2052FG-10A", "8"),
        ],
        [Inches(0.9), Inches(0.45), Inches(1.35), Inches(0.8), Inches(2.35), Inches(0.45)],
    )

    doc.add_heading("4. กรอกรายละเอียดและสร้าง Job", level=2)
    add_image(doc, "03-admin-create-job.png", "ภาพหน้าสร้าง Job จาก PO")
    add_steps(
        doc,
        [
            "ตรวจกรอบ รายการที่เลือกจาก PO รอจัดส่ง ว่ามี PO ครบตามที่เลือก",
            "กรอก ชื่อห้อง Job เช่น รอบเช้า บางซื่อ-ลาดพร้าว / ส่งของร้าน A",
            "กรอก รถขนส่ง เช่น 6W-4382 หรือทะเบียนรถจริง",
            "กรอก คนขับ เช่น สมชาย หรือชื่อคนขับที่รับผิดชอบ",
            "ตรวจ ต้นทาง ค่าเริ่มต้นคือ DC Bangna; ช่อง GPS ต้นทางจะดึงจากมือถือของผู้เริ่มงานใน Driver Room เท่านั้น",
            "เพิ่มปลายทาง กรอกชื่อหรือโลเคชันปลายทาง แล้วติ๊ก PO ที่ต้องส่งไปจุดนั้น",
            "กด สร้าง Job จริงจากรายการที่เลือก ระบบจะพาไปหน้า Monitor ของ Job นั้นทันที",
        ],
    )
    add_table(
        doc,
        ["ช่องกรอก", "ตัวอย่างข้อมูลจำลอง"],
        [
            ("ชื่อห้อง Job", "รอบเช้า บางซื่อ-ลาดพร้าว"),
            ("รถขนส่ง", "6W-4382"),
            ("คนขับ", "สมชาย"),
            ("ต้นทาง", "DC Bangna"),
            ("ปลายทาง", "ร้าน A / Mixed Cement, Mortar&Decorative Brand MG"),
            ("หมายเหตุ", "ส่งก่อน 10:30 น. ติดต่อคลังปลายทางก่อนลงของ"),
        ],
        [Inches(1.7), Inches(4.7)],
    )

    doc.add_heading("5. Monitor Realtime", level=2)
    add_image(doc, "04-admin-monitor.png", "ภาพหน้า Monitor Realtime สำหรับดูสถานะงาน")
    add_steps(
        doc,
        [
            "เปิดเมนู Monitor Realtime หรือกดปุ่ม Monitor จากหน้ารายการ Job",
            "เลือก Job ที่ต้องการดู หากเข้าจากปุ่ม Monitor ระบบจะเปิด Job นั้นให้ทันที",
            "ดูสถานะหลัก เช่น Status, Route, Alerts, จำนวนต้องสแกน, Loaded และ Delivered",
            "ตรวจแผนส่งตาม Location / PO ว่าปลายทางและ PO ถูกต้อง",
            "ถ้าต้องการให้คนขับเข้าใช้งาน กด เปิด Driver Room หรือ แสดง QR สำหรับคนขับ",
        ],
    )
    doc.add_paragraph("ข้อมูลตัวอย่างจาก Job ที่ใช้ในคู่มือ:")
    add_table(
        doc,
        ["รายการ", "ค่า"],
        [
            ("Job ID", "JOB-20260518-152703-000012"),
            ("ห้อง Job", "frt"),
            ("Status", "ready"),
            ("Route", "1 Locations"),
            ("PO", "8030488988"),
            ("รหัสวัสดุ", "80215392"),
            ("สินค้า", "ตลับเมตรโลโก้ เสือ ยี่ห้อ Stanley"),
            ("จำนวนในไฟล์", "5,040"),
            ("ต้องสแกน", "1"),
        ],
        [Inches(1.65), Inches(4.75)],
    )

    doc.add_heading("5.1 กติกาล็อกต้นทางและเปิดกรณีพิเศษโดย Admin", level=2)
    add_bullets(
        doc,
        [
            "เมื่อคนขับเช็กอินต้นทางและสแกนขึ้นรถครบตามจำนวน ต้องสแกน แล้ว ระบบควรถือว่าต้นทางปิดงานทันที",
            "หลังปิดต้นทาง ไม่ควรให้คนขับกดเช็กอินต้นทางซ้ำหรือเปลี่ยนพิกัดต้นทางเอง เพื่อป้องกันการเช็กอินผิดที่หลังรถออกจากคลัง",
            "สถานะที่ Admin ควรเห็นคือ ขึ้นรถครบ / ต้นทางปิด / พร้อมไปปลายทาง",
            "หากมีเหตุผิดปกติ เช่น สแกนขาด, สแกนผิด, ต้องกลับไปโหลดเพิ่ม หรือ GPS ผิด ให้ Admin เป็นคนตรวจ Monitor และเปิดสิทธิ์กรณีพิเศษเท่านั้น",
            "แนวคิดเดียวกันใช้กับปลายทาง: คนขับทำตามขั้นตอนปกติ ส่วนการเปิดปลายทางกรณีพิเศษหรือปลดล็อกงานต้องเป็นหน้าที่ Admin",
        ],
    )
    add_table(
        doc,
        ["กรณีพิเศษ", "ผู้ทำรายการ", "วิธีจัดการ"],
        [
            ("คนขับสแกนขึ้นรถครบแล้ว แต่กดเช็กอินต้นทางผิดที่ภายหลัง", "Admin", "ไม่ให้คนขับแก้เอง ให้ Admin ตรวจ Monitor และใช้ข้อมูลต้นทางที่ปิดไว้เป็นหลัก"),
            ("ต้องกลับไปโหลดสินค้าเพิ่ม", "Admin", "Admin เปิดสิทธิ์/ปลดล็อกต้นทางชั่วคราว แล้วให้คนขับสแกนเพิ่มตามรายการจริง"),
            ("ถึงปลายทางแต่ GPS ไม่ตรงหรืออยู่นอกรัศมี", "Admin", "Admin ใช้ปุ่มเปิดปลายทางกรณีพิเศษ หลังตรวจสอบกับคนขับแล้ว"),
            ("สแกนผิดสินค้า หรือสแกนรหัสที่ไม่อยู่ใน Job", "Admin", "ดู Alert Queue และตัดสินใจว่าจะให้สแกนใหม่ ลบรายการผิด หรือเปิดกรณีพิเศษ"),
            ("มือถือคนขับเสีย/แบตหมด ต้องเปลี่ยนเครื่อง", "Admin", "Admin ตรวจสถานะล่าสุด แล้วส่ง QR/ลิงก์ให้เครื่องใหม่ พร้อมกำชับว่าห้ามเช็กอินต้นทางซ้ำถ้าต้นทางปิดแล้ว"),
        ],
        [Inches(2.2), Inches(1.0), Inches(3.2)],
    )

    doc.add_heading("6. ส่งลิงก์หรือ QR ให้คนขับ", level=2)
    add_steps(
        doc,
        [
            "ในหน้า Monitor หรือหน้ารายการ Job กด แสดง QR สำหรับคนขับ",
            "ให้คนขับสแกน QR ด้วยมือถือเครื่องที่จะใช้เช็กอิน GPS และสแกนสินค้า",
            "หากสแกนไม่ได้ ให้ส่งลิงก์ Driver Room โดยตรง เช่น https://project-stock-qr.vercel.app/driver-room?jobId=JOB-20260518-152703-000012",
            "ย้ำว่าการเช็กอิน GPS ควรทำจากมือถือของคนขับ ไม่ใช่เครื่องของแอดมิน",
        ],
    )

    doc.add_heading("7. ดูประวัติงาน", level=2)
    add_image(doc, "06-admin-history.png", "ภาพหน้าประวัติงาน สำหรับค้นหางานที่ปิดจบแล้ว")
    add_steps(
        doc,
        [
            "กดเมนู ประวัติงาน",
            "ค้นหาด้วย Job ID, คนขับ, รถ, PO หรือรหัสวัสดุ",
            "กรองช่วงวันที่ปิดงานหากต้องการดูย้อนหลัง",
            "เลือกงานจากรายการเพื่อเปิดรายละเอียดงานที่ปิดจบแล้ว",
        ],
    )

    doc.add_heading("คู่มือฝั่งคนขับ", level=1)

    doc.add_heading("หลักการหน้าคนขับสำหรับผู้สูงอายุ", level=2)
    add_bullets(
        doc,
        [
            "ให้คนขับทำทีละขั้นตอนตามตัวเลข 1, 2, 3 บนหน้าจอ ไม่ต้องจำเมนูเยอะ",
            "ปุ่มหลักที่ต้องกดควรเป็นปุ่มใหญ่ เช่น เช็กอินต้นทางเพื่อเริ่มสแกน, เปิดกล้อง, บันทึก และเช็กอินปลายทาง",
            "ข้อความบนหน้าจอควรใช้คำสั้น ๆ เช่น รอ GPS, โหลดครบแล้ว, ส่งครบแล้ว, ยังสแกนไม่ครบ",
            "ถ้าขึ้นข้อความเตือน สีเหลือง หรือสแกนไม่ผ่าน ให้หยุดและโทรหา Admin ไม่ต้องกดซ้ำหลายครั้ง",
            "หลังสแกนขึ้นรถครบแล้ว คนขับไม่ต้องกลับไปกดเช็กอินต้นทางอีก ระบบจะถือว่าปิดต้นทางแล้ว",
        ],
    )

    doc.add_heading("1. เปิด Driver Room", level=2)
    add_image(doc, "07-driver-room-desktop.png", "ภาพหน้า Driver Room เมื่อเปิดจากลิงก์ Job")
    add_steps(
        doc,
        [
            "เปิดลิงก์ที่ Admin ส่งให้ หรือสแกน QR ของ Job",
            "ตรวจหัวหน้าจอว่าห้อง Job, Job ID, รถ และชื่อคนขับถูกต้อง",
            "ดูตัวเลข ต้องสแกน, ขึ้นรถแล้ว และส่งแล้ว เพื่อรู้จำนวนงานที่ต้องทำ",
        ],
    )
    doc.add_paragraph("ตัวอย่างข้อมูลจำลองที่ใช้สอนคนขับ:")
    add_table(
        doc,
        ["รายการ", "ตัวอย่าง"],
        [
            ("ห้อง Job", "A"),
            ("Job ID", "JOB-20260518-085654-000014"),
            ("รถ", "1234"),
            ("คนขับ", "b"),
            ("ต้นทาง", "บางซ่อน"),
            ("ปลายทาง", "ร้าน A / โซนรับสินค้า 1"),
            ("ต้องสแกน", "2 ชิ้นหรือ 2 รอบสแกน"),
            ("รหัสตัวอย่างสำหรับสแกน", "8030492670-1 และ 8030492670-2"),
        ],
        [Inches(1.8), Inches(4.6)],
    )

    doc.add_heading("2. เช็กอินต้นทาง", level=2)
    add_steps(
        doc,
        [
            "อยู่ที่ต้นทางจริง เช่น DC Bangna",
            "กด เช็กอินต้นทาง",
            "อนุญาตให้ browser เข้าถึงตำแหน่ง GPS หากมือถือถามสิทธิ์",
            "รอจนระบบบันทึกตำแหน่งสำเร็จ ข้อความสถานะจะเปลี่ยนจาก ยังไม่เช็กอินต้นทาง เป็นเช็กอินแล้ว",
            "หลังเช็กอินสำเร็จ ระบบจะแสดงปุ่ม เปิดกล้อง และช่องกรอกรหัสสำหรับสแกนสินค้า",
        ],
    )

    doc.add_heading("3. สแกนขึ้นรถ", level=2)
    add_steps(
        doc,
        [
            "หลังเช็กอินต้นทาง ให้เลือกโหมด ขึ้นรถ หรือ โหลดสินค้าที่คลัง",
            "กด เปิดกล้อง หากมือถือถามสิทธิ์ Camera ให้กดอนุญาต",
            "หันกล้องไปที่ QR Code หรือ Barcode ของสินค้า ให้รหัสอยู่กลางกรอบและนิ่งประมาณ 1-2 วินาที",
            "เมื่อสแกนสำเร็จ ระบบจะบันทึก 1 รายการและตัวเลข ขึ้นรถแล้ว จะเพิ่มขึ้น เช่น จาก 0 เป็น 1",
            "สแกนชิ้นถัดไปด้วยวิธีเดิม จนตัวเลข ขึ้นรถแล้ว เท่ากับ ต้องสแกน เช่น ต้องสแกน 2 และขึ้นรถแล้ว 2",
            "เมื่อขึ้นรถครบแล้ว ให้ถือว่าต้นทางปิดทันที คนขับไม่ต้องกดเช็กอินต้นทางซ้ำ และไม่ควรเปิดหน้านี้เพื่อเช็กอินต้นทางจากที่อื่น",
            "ถ้าสแกนไม่ติด ให้เช็ดเลนส์ เพิ่มแสง ขยับมือถือให้ห่างจากโค้ดเล็กน้อย แล้วลองใหม่",
        ],
    )
    doc.add_paragraph("ตัวอย่างการสแกนขึ้นรถ:")
    add_table(
        doc,
        ["รอบ", "สิ่งที่คนขับทำ", "รหัสตัวอย่าง", "ผลที่ควรเห็น"],
        [
            ("1", "สแกน QR ชิ้นแรก", "8030492670-1", "ขึ้นรถแล้ว 0 -> 1"),
            ("2", "สแกน QR ชิ้นที่สอง", "8030492670-2", "ขึ้นรถแล้ว 1 -> 2 และสถานะโหลดครบ / ปิดต้นทาง"),
            ("กรณีไม่มี QR", "กรอกรหัสในช่อง เลข PO / Barcode / QR / registry key แล้วกด บันทึก", "8030492670-1", "ระบบนับเหมือนการสแกนด้วยกล้อง"),
        ],
        [Inches(0.55), Inches(2.25), Inches(1.5), Inches(2.1)],
    )

    doc.add_heading("4. กรอกรหัสแทนการสแกนด้วยกล้อง", level=2)
    add_steps(
        doc,
        [
            "ใช้วิธีนี้เมื่อกล้องเปิดไม่ได้, QR เสีย, Barcode เลือน หรือหน้างานแสงน้อย",
            "เลือกโหมดให้ถูกก่อนเสมอ ถ้ากำลังโหลดขึ้นรถให้เลือก ขึ้นรถ ถ้ากำลังส่งของให้เลือก ส่งปลายทาง",
            "แตะช่อง เลข PO / Barcode / QR / registry key",
            "พิมพ์รหัสตามฉลากหรือใบงาน เช่น 8030492670-1 หรือ 8030492670-2",
            "กด บันทึก แล้วดูว่าตัวเลข ขึ้นรถแล้ว หรือ ส่งแล้ว เพิ่มขึ้นตามโหมดที่เลือก",
        ],
    )

    doc.add_heading("5. เช็กอินและสแกนส่งปลายทาง", level=2)
    add_steps(
        doc,
        [
            "เมื่อถึงปลายทาง ให้กด เช็กอินปลายทาง",
            "อนุญาตตำแหน่ง GPS หากมือถือถามสิทธิ์ และรอจนระบบบันทึกพิกัดปลายทางสำเร็จ",
            "เลือกโหมด ส่งปลายทาง / ส่งของให้ลูกค้า",
            "สแกน QR Code หรือ Barcode ของสินค้าที่ส่งจริงทีละชิ้น หรือกรอกรหัสแทนหากกล้องใช้ไม่ได้",
            "ตรวจตัวเลข ส่งแล้ว ว่าเพิ่มขึ้น เช่น จาก 0 เป็น 1 แล้วเป็น 2",
            "เมื่อส่งครบ ระบบจะแสดงว่างานปลายทางนั้นส่งครบหรือสถานะเปลี่ยนเป็นส่งแล้ว",
            "หากมีหลายปลายทาง ให้ทำซ้ำตั้งแต่เช็กอินปลายทางตามลำดับปลายทางที่ระบบแสดง",
        ],
    )
    doc.add_paragraph("ตัวอย่างการสแกนส่งปลายทาง:")
    add_table(
        doc,
        ["รอบ", "สิ่งที่คนขับทำ", "รหัสตัวอย่าง", "ผลที่ควรเห็น"],
        [
            ("1", "เลือกโหมดส่งปลายทางแล้วสแกนชิ้นแรก", "8030492670-1", "ส่งแล้ว 0 -> 1"),
            ("2", "สแกนชิ้นที่สอง", "8030492670-2", "ส่งแล้ว 1 -> 2"),
            ("จบงาน", "ตรวจยอด ต้องสแกน 2 / ขึ้นรถแล้ว 2 / ส่งแล้ว 2", "-", "งานปลายทางครบ ไม่มีรายการค้าง"),
        ],
        [Inches(0.55), Inches(2.45), Inches(1.4), Inches(2.0)],
    )

    doc.add_heading("6. ถ้าสแกนผิดหรือสแกนซ้ำ", level=2)
    add_table(
        doc,
        ["เหตุการณ์", "สิ่งที่ควรทำ"],
        [
            ("สแกนรหัสที่ไม่ได้อยู่ใน Job", "หยุดสแกนรายการนั้น แจ้ง Admin ให้ดู Alert Queue ในหน้า Monitor"),
            ("สแกนซ้ำรหัสเดิม", "ตรวจยอดบนหน้าจอ ถ้าระบบแจ้งเตือนหรือยอดไม่เพิ่ม ให้แจ้ง Admin ไม่ต้องสแกนซ้ำหลายครั้ง"),
            ("เลือกโหมดผิด เช่น ส่งปลายทางตอนกำลังขึ้นรถ", "เปลี่ยนโหมดให้ถูก แล้วแจ้ง Admin ตรวจประวัติการสแกนใน Monitor"),
            ("ยอดขึ้นรถครบแต่ส่งแล้วยังไม่ครบ", "ยังปิดงานไม่ได้ ให้ตรวจว่ายังมีสินค้าที่ปลายทางต้องสแกนส่งอีกหรือไม่"),
            ("ขึ้นรถครบแล้ว แต่ต้องแก้ต้นทางหรือโหลดเพิ่ม", "คนขับไม่เปิดเอง ให้โทรหา Admin เพื่อเปิดกรณีพิเศษ"),
            ("ถึงปลายทางแล้วระบบไม่ให้เช็กอิน", "อย่ากดมั่ว ให้โทรหา Admin ตรวจ GPS/รัศมี และเปิดปลายทางกรณีพิเศษถ้าจำเป็น"),
        ],
        [Inches(2.1), Inches(4.3)],
    )

    doc.add_heading("7. มุมมองบนมือถือ", level=2)
    add_image(doc, "08-driver-room-mobile.png", "ภาพตัวอย่าง Driver Room บนหน้าจอมือถือ")
    add_bullets(
        doc,
        [
            "ใช้มือถือเครื่องเดียวกันตลอดงาน เพื่อให้ GPS ต้นทางและปลายทางตรงกับผู้ปฏิบัติงานจริง",
            "ถ้า browser ขอสิทธิ์กล้องหรือพิกัด ให้กดอนุญาต",
            "ถ้าสแกนแล้วไม่ผ่าน ให้ตรวจว่าอยู่โหมดถูกต้อง: ขึ้นรถ หรือ ส่งปลายทาง",
            "ถ้ามี alert หรือสแกนผิดสินค้า ให้แจ้ง Admin ให้ตรวจหน้า Monitor",
        ],
    )

    doc.add_heading("ข้อควรระวังและการแก้ปัญหา", level=1)
    add_table(
        doc,
        ["อาการ", "วิธีตรวจ / วิธีแก้"],
        [
            ("ปุ่มสร้าง Job กดไม่ได้", "ยังไม่ได้เลือก PO หรือรายการที่เลือกไม่มีในคิว ให้กลับไปหน้า PO รอจัดส่งแล้วติ๊ก checkbox ก่อน"),
            ("คนขับเปิด Driver Room แล้วไม่เห็นงาน", "ตรวจว่าใช้ลิงก์ที่มี jobId ถูกต้อง หรือให้สแกน QR ใหม่จากหน้า Monitor"),
            ("เช็กอิน GPS ไม่ได้", "เปิดสิทธิ์ Location ในมือถือและ browser แล้วลองกดเช็กอินใหม่จากจุดจริง"),
            ("เปิดกล้องไม่ได้", "อนุญาตสิทธิ์ Camera, ปิด browser แล้วเปิดลิงก์ใหม่ หรือกรอกรหัสในช่อง manual แทน"),
            ("ยอด Loaded / Delivered ไม่ตรง", "ให้ Admin ตรวจ Monitor เทียบกับรายการในห้อง และดู Alert Queue"),
            ("ต้องเปิดปลายทางกรณีพิเศษ", "Admin ใช้ปุ่ม Admin เปิดปลายทางกรณีพิเศษ ในหน้า Monitor เฉพาะกรณีจำเป็น"),
            ("คนขับสแกนขึ้นรถครบแล้วแต่ไปเช็กอินต้นทางผิดที่", "ให้ยึดต้นทางที่ปิดตอนโหลดครบ ห้ามให้คนขับแก้เอง Admin ตรวจและเปิดกรณีพิเศษเท่านั้น"),
            ("คนขับสูงอายุกดไม่ถูก", "ให้ทำตามเลขขั้นตอนบนหน้าจอ 1 เช็กอินต้นทาง, 2 สแกนขึ้นรถ, 3 สแกนสินค้า และโทรหา Admin เมื่อมีข้อความเตือน"),
        ],
        [Inches(2.0), Inches(4.4)],
    )

    doc.add_heading("Checklist ก่อนใช้งานจริง", level=1)
    add_bullets(
        doc,
        [
            "ไฟล์ PO ถูกนำเข้าและตรวจข้อมูลหลักครบถ้วน",
            "Admin เลือก PO ถูกชุดก่อนสร้าง Job",
            "ชื่อห้อง Job, รถ, คนขับ, ต้นทาง และปลายทางถูกต้อง",
            "คนขับได้รับ QR หรือลิงก์ Driver Room ของ Job ที่ถูกต้อง",
            "มือถือคนขับเปิดสิทธิ์ GPS และกล้องเรียบร้อย",
            "Admin เปิดหน้า Monitor เพื่อตรวจสถานะและ alert ระหว่างงาน",
            "ตกลงร่วมกันว่าหลังสแกนขึ้นรถครบ ต้นทางถือว่าปิดแล้ว คนขับไม่เช็กอินต้นทางซ้ำ",
            "กำหนดคนรับผิดชอบฝั่ง Admin สำหรับเปิดกรณีพิเศษ เช่น เปิดปลายทาง/ปลดล็อกต้นทาง",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Project Stock QR Manual | 18 May 2026")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
